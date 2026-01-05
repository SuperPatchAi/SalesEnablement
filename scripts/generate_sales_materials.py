#!/usr/bin/env python3
"""
Super Patch Sales Materials Generator

Generates comprehensive word track documents (.docx) and visual infographics
for each Super Patch product using Google Gemini and Imagen 3.

Usage:
    python generate_sales_materials.py                    # Generate all materials
    python generate_sales_materials.py --product freedom  # Generate for specific product
    python generate_sales_materials.py --docx-only        # Only generate DOCX files
    python generate_sales_materials.py --images-only      # Only generate infographics
"""

import os
import json
import sys
import argparse
from pathlib import Path
from datetime import datetime

# Document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

# Google Gemini
from google import genai
from google.genai import types

# Image handling
from PIL import Image
import io
import base64


# Configuration
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
OUTPUT_DIR = Path('/workspace/sales_materials')
PRODUCTS_FILE = Path('/workspace/products/superpatch_products.json')
FRAMEWORKS_DIR = Path('/workspace/frameworks')
PRODUCT_REF_FILE = Path('/workspace/products/SuperPatch_Product_Reference.md')


def load_products():
    """Load product data from JSON file."""
    with open(PRODUCTS_FILE, 'r') as f:
        return json.load(f)


def load_frameworks():
    """Load all framework content from markdown files."""
    frameworks = {}
    for md_file in FRAMEWORKS_DIR.glob('*.md'):
        with open(md_file, 'r') as f:
            frameworks[md_file.stem] = f.read()
    return frameworks


def load_product_reference():
    """Load the detailed product reference."""
    with open(PRODUCT_REF_FILE, 'r') as f:
        return f.read()


def create_gemini_client():
    """Initialize the Gemini client."""
    return genai.Client(api_key=GEMINI_API_KEY)


def generate_word_track_content(client, product_key: str, product_data: dict, 
                                 frameworks: dict, product_reference: str) -> dict:
    """
    Use Gemini to generate comprehensive word track content for a product.
    """
    
    # Select most relevant frameworks
    relevant_frameworks = []
    for name in ['objection_handling', 'closing_techniques', 'sales_psychology_and_building_rapport', 
                 'direct_sales_and_network_marketing', 'follow-up_strategies']:
        if name in frameworks:
            # Get first 8000 chars to stay within limits
            relevant_frameworks.append(f"## {name.replace('_', ' ').title()}\n{frameworks[name][:8000]}")
    
    frameworks_text = "\n\n---\n\n".join(relevant_frameworks)
    
    prompt = f"""You are a sales training expert creating a comprehensive Word Track document for a direct sales team selling Super Patch wellness products.

## PRODUCT INFORMATION
Product Name: {product_data['name']}
Category: {product_data['category']}
Tagline: "{product_data['tagline']}"
Benefits: {', '.join(product_data['benefits'])}
Target Audience: {product_data['target_audience']}

## ABOUT SUPER PATCH
Super Patch uses Vibrotactile Technology (VTT) - specialized ridge patterns that interact with the skin's mechanoreceptors to trigger the body's own neural responses. It's 100% drug-free, all-natural, and non-invasive.

## SALES FRAMEWORKS & TECHNIQUES (Use these proven methods)
{frameworks_text[:20000]}

---

## YOUR TASK
Create a COMPLETE Word Track document for the {product_data['name']} patch that a salesperson can use as their daily reference. Include:

### 1. PRODUCT OVERVIEW (2-3 paragraphs)
- What the product does
- Who it's for
- Why it's unique

### 2. IDEAL CUSTOMER PROFILE
- Demographics
- Psychographics  
- Pain points they experience
- What they've tried before

### 3. OPENING SCRIPTS (5 different openers)
Write 5 complete, natural conversation starters for different scenarios:
- Cold approach (stranger)
- Warm introduction (friend/family)
- Social media DM
- Referral introduction
- Event/party approach

### 4. DISCOVERY QUESTIONS (10 questions)
Questions to uncover needs and pain points. Include:
- Opening questions
- Pain point questions
- Impact questions
- Solution questions

### 5. PRODUCT PRESENTATION SCRIPT
A complete 2-minute presentation script using Problem-Agitate-Solve framework

### 6. OBJECTION HANDLING (8 objections with scripts)
Complete response scripts for each objection:
1. "It's too expensive"
2. "I need to think about it"
3. "I need to talk to my spouse"
4. "Does it really work?"
5. "I've tried patches/products before"
6. "I'm not interested"
7. "I don't have time"
8. "How is this different from [competitor]?"

For each: Give the exact word-for-word response using techniques from the frameworks above.

### 7. CLOSING SCRIPTS (5 closes)
Different closing techniques with exact scripts:
- The Assumptive Close
- The Alternative Close  
- The Urgency Close
- The Summary Close
- The Referral Close

### 8. FOLLOW-UP SEQUENCES
- Day 1, 3, 7, 14 follow-up messages
- Voicemail scripts
- Text message templates

### 9. TESTIMONIAL PROMPTS
Questions to ask satisfied customers to get powerful testimonials

### 10. QUICK REFERENCE CARD
A concise "cheat sheet" with:
- 3 key benefits
- 3 best discovery questions
- Top 3 objection responses
- 2 best closing lines

Make ALL scripts sound natural, conversational, and authentic. Use the sales techniques from the frameworks provided. Be specific and actionable.
"""

    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                temperature=0.7,
                max_output_tokens=8000,
            )
        )
        
        return {
            'success': True,
            'content': response.text,
            'product': product_data['name']
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'product': product_data['name']
        }


def create_docx(content: str, product_name: str, product_data: dict, output_path: Path):
    """
    Create a professionally formatted Word document from the generated content.
    """
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.author = 'Super Patch AI Sales Team'
    core_properties.title = f'{product_name} Patch - Sales Word Track'
    
    # Add title
    title = doc.add_heading(f'{product_name} Patch', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Complete Sales Word Track Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add tagline
    tagline = doc.add_paragraph(f'"{product_data["tagline"]}"')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.runs[0]
    tagline_run.font.italic = True
    tagline_run.font.size = Pt(14)
    
    # Add horizontal line
    doc.add_paragraph('_' * 60)
    
    # Add product info box
    info_para = doc.add_paragraph()
    info_para.add_run('Category: ').bold = True
    info_para.add_run(f'{product_data["category"]}  |  ')
    info_para.add_run('Benefits: ').bold = True
    info_para.add_run(', '.join(product_data['benefits']))
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Parse and add content sections
    lines = content.split('\n')
    current_para = None
    
    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Main section headers (### 1., ### 2., etc.)
        if line.startswith('### ') or line.startswith('## '):
            header_text = line.lstrip('#').strip()
            # Remove numbering like "1. " from start
            if header_text and header_text[0].isdigit() and '. ' in header_text[:4]:
                header_text = header_text.split('. ', 1)[1] if '. ' in header_text else header_text
            doc.add_heading(header_text, level=1 if line.startswith('## ') else 2)
            current_para = None
            
        # Sub-headers (####)
        elif line.startswith('#### '):
            header_text = line.lstrip('#').strip()
            doc.add_heading(header_text, level=3)
            current_para = None
            
        # Bold items (** at start)
        elif line.startswith('**') and '**' in line[2:]:
            para = doc.add_paragraph()
            # Parse bold sections
            parts = line.split('**')
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
            current_para = para
            
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:]
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            current_para = para
            
        # Numbered items
        elif line[0].isdigit() and ('. ' in line[:4] or ') ' in line[:4]):
            # Extract the text after the number
            if '. ' in line[:4]:
                text = line.split('. ', 1)[1] if '. ' in line else line
            else:
                text = line.split(') ', 1)[1] if ') ' in line else line
            para = doc.add_paragraph(text, style='List Number')
            current_para = para
            
        # Quote blocks (scripts)
        elif line.startswith('"') or line.startswith('>'):
            text = line.lstrip('> ').strip('"')
            para = doc.add_paragraph()
            run = para.add_run(f'"{text}"')
            run.font.italic = True
            para.paragraph_format.left_indent = Inches(0.5)
            current_para = para
            
        # Regular text
        else:
            para = doc.add_paragraph(line)
            current_para = para
    
    # Add footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    footer = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")} | Super Patch Sales Enablement')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save(output_path)
    print(f"✅ Created: {output_path}")


def generate_infographic(client, product_name: str, product_data: dict, output_path: Path) -> bool:
    """
    Generate a visual infographic using Imagen 4.
    """
    
    prompt = f"""Create a clean, professional sales infographic for a wellness product called "{product_name} Patch".

Design requirements:
- Modern, clean design with soft gradient background (light blue to white)
- Professional color scheme (blues, greens, white)
- Clear visual hierarchy

Content to include:
- Title: "{product_name}" at the top
- Tagline: "{product_data['tagline']}"
- 3 key benefits as icons with text:
  1. {product_data['benefits'][0]}
  2. {product_data['benefits'][1]}
  3. {product_data['benefits'][2] if len(product_data['benefits']) > 2 else 'Easy to use'}
- A simple 3-step process at bottom: "Apply → Wear → Feel the Difference"
- "Drug-Free | Natural | Non-Invasive" footer text

Style: Corporate wellness, medical, trustworthy, modern flat design, infographic style
NO: text errors, blurry text, photorealistic people, complex scenes"""

    try:
        response = client.models.generate_images(
            model='imagen-4.0-generate-001',
            prompt=prompt,
            config=types.GenerateImagesConfig(
                number_of_images=1,
                aspect_ratio='3:4',  # Portrait for document insert
                safety_filter_level='BLOCK_LOW_AND_ABOVE',
            )
        )
        
        if response.generated_images:
            image_data = response.generated_images[0].image.image_bytes
            
            # Save the image
            with open(output_path, 'wb') as f:
                f.write(image_data)
            
            print(f"✅ Created infographic: {output_path}")
            return True
        else:
            print(f"⚠️ No image generated for {product_name}")
            return False
            
    except Exception as e:
        print(f"❌ Error generating infographic for {product_name}: {e}")
        return False


def generate_process_infographic(client, product_name: str, product_data: dict, output_path: Path) -> bool:
    """
    Generate a step-by-step sales process infographic.
    """
    
    prompt = f"""Create a step-by-step sales process infographic for selling "{product_name} Patch" wellness product.

Design requirements:
- Horizontal flow diagram with 5 connected steps
- Clean, modern corporate design
- Use arrows or flowing lines connecting each step
- Professional blue/green color scheme

The 5 steps should be:
Step 1: "CONNECT" - Start conversation, build rapport
Step 2: "DISCOVER" - Ask about their {product_data['category'].lower()} challenges  
Step 3: "PRESENT" - Explain how {product_name} helps with {product_data['benefits'][0].lower()}
Step 4: "HANDLE OBJECTIONS" - Address concerns with confidence
Step 5: "CLOSE" - Ask for the sale, get commitment

Each step should have:
- A number (1-5)
- A one-word title
- A short description (5-7 words)
- Simple icon representation

Style: Corporate training material, infographic, modern flat design, flowchart
Make it easy to follow from left to right
NO: photorealistic images, complex backgrounds, blurry text"""

    try:
        response = client.models.generate_images(
            model='imagen-4.0-generate-001',
            prompt=prompt,
            config=types.GenerateImagesConfig(
                number_of_images=1,
                aspect_ratio='16:9',  # Landscape for process flow
                safety_filter_level='BLOCK_LOW_AND_ABOVE',
            )
        )
        
        if response.generated_images:
            image_data = response.generated_images[0].image.image_bytes
            
            with open(output_path, 'wb') as f:
                f.write(image_data)
            
            print(f"✅ Created process infographic: {output_path}")
            return True
        else:
            print(f"⚠️ No process image generated for {product_name}")
            return False
            
    except Exception as e:
        print(f"❌ Error generating process infographic for {product_name}: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(description='Generate Super Patch sales materials')
    parser.add_argument('--product', type=str, help='Generate for specific product (e.g., freedom, liberty)')
    parser.add_argument('--docx-only', action='store_true', help='Only generate DOCX files')
    parser.add_argument('--images-only', action='store_true', help='Only generate infographics')
    parser.add_argument('--list', action='store_true', help='List available products')
    args = parser.parse_args()
    
    # Load data
    print("📂 Loading product data and frameworks...")
    products = load_products()
    frameworks = load_frameworks()
    product_reference = load_product_reference()
    
    if args.list:
        print("\n📋 Available Products:")
        for key, data in products.items():
            print(f"  - {key}: {data['name']} ({data['category']})")
        return
    
    # Create output directories
    OUTPUT_DIR.mkdir(exist_ok=True)
    (OUTPUT_DIR / 'docx').mkdir(exist_ok=True)
    (OUTPUT_DIR / 'infographics').mkdir(exist_ok=True)
    
    # Initialize Gemini client
    print("🤖 Initializing Gemini client...")
    client = create_gemini_client()
    
    # Determine which products to process
    if args.product:
        if args.product.lower() not in products:
            print(f"❌ Product '{args.product}' not found. Use --list to see available products.")
            return
        products_to_process = {args.product.lower(): products[args.product.lower()]}
    else:
        products_to_process = products
    
    print(f"\n🚀 Processing {len(products_to_process)} product(s)...\n")
    
    # Process each product
    results = {'docx': [], 'infographics': []}
    
    for product_key, product_data in products_to_process.items():
        print(f"\n{'='*50}")
        print(f"📦 Processing: {product_data['name']} Patch")
        print(f"{'='*50}")
        
        # Generate DOCX
        if not args.images_only:
            print(f"\n📝 Generating word track document...")
            result = generate_word_track_content(
                client, product_key, product_data, frameworks, product_reference
            )
            
            if result['success']:
                docx_path = OUTPUT_DIR / 'docx' / f"{product_data['name']}_WordTrack.docx"
                create_docx(result['content'], product_data['name'], product_data, docx_path)
                results['docx'].append(product_data['name'])
                
                # Also save markdown version
                md_path = OUTPUT_DIR / 'docx' / f"{product_data['name']}_WordTrack.md"
                with open(md_path, 'w') as f:
                    f.write(f"# {product_data['name']} Patch - Sales Word Track\n\n")
                    f.write(f"*{product_data['tagline']}*\n\n")
                    f.write(f"**Category:** {product_data['category']} | ")
                    f.write(f"**Benefits:** {', '.join(product_data['benefits'])}\n\n")
                    f.write("---\n\n")
                    f.write(result['content'])
                print(f"✅ Created: {md_path}")
            else:
                print(f"❌ Failed to generate content: {result['error']}")
        
        # Generate infographics
        if not args.docx_only:
            print(f"\n🎨 Generating infographics...")
            
            # Product overview infographic
            info_path = OUTPUT_DIR / 'infographics' / f"{product_data['name']}_Overview.png"
            if generate_infographic(client, product_data['name'], product_data, info_path):
                results['infographics'].append(f"{product_data['name']}_Overview")
            
            # Sales process infographic
            process_path = OUTPUT_DIR / 'infographics' / f"{product_data['name']}_SalesProcess.png"
            if generate_process_infographic(client, product_data['name'], product_data, process_path):
                results['infographics'].append(f"{product_data['name']}_SalesProcess")
    
    # Summary
    print(f"\n{'='*50}")
    print("📊 GENERATION COMPLETE")
    print(f"{'='*50}")
    print(f"\n✅ Word Track Documents: {len(results['docx'])}")
    for name in results['docx']:
        print(f"   - {name}_WordTrack.docx")
    
    print(f"\n✅ Infographics: {len(results['infographics'])}")
    for name in results['infographics']:
        print(f"   - {name}.png")
    
    print(f"\n📁 Output directory: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
